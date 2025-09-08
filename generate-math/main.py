from docx import Document
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.shared import Pt
from fractions import Fraction
from random import randint, choice
import qrcode
from datetime import datetime
import os
# 创建变量
symbols = [" + ", " - ", " × ", " ÷ "]
question_types = ["混合", "加法", "减法", "乘法", "除法"]
# 生成随机分数
def Random_Fraction():
    while True:
        f = Fraction(randint(1,10), randint(1,10))
        if f < 2 and f != 1:
            return f
# 生成随机问题
def Random_Question(question_type):
    global symbols
    f1 = Random_Fraction()
    if question_type:
        symbol = symbols[question_type-1]
    else:
        symbol = choice(symbols)
    f2 = Random_Fraction()
    match symbol:
        case " + ":
            ans = f1 + f2
        case " - ":
            ans = f1 - f2
        case " × ":
            ans = f1 * f2
        case " ÷ ":
            ans = f1 / f2
    #ans = eval("f1" + symbol + "f2")
    return [f1 , symbol , f2], ans
# 生成试卷
def generate_paper(question_type, quantity, folder):
    global question_types
    questions = []
    answers = []
    # 新建文档
    doc_out = Document()
    # 获取文档的第一个节（默认只有一个节）
    section = doc_out.sections[0]
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
    # 设置页边距（单位为厘米）
    section.top_margin = Cm(1.27) # 上边距
    section.bottom_margin = Cm(1.27) # 下边距
    section.left_margin = Cm(1.27) # 左边距
    section.right_margin = Cm(1.27) # 右边距
    paragraph = doc_out.add_paragraph()
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(text=str(quantity) + f"道分数{question_types[question_type]}四则运算")
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(24)
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    paragraph = doc_out.add_paragraph()
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(text="班级：____________姓名：____________")
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 开始生成
    for x in range(quantity):
        while True:
            question, answer = Random_Question(question_type)
            if len(questions) > 0:
                if questions.count(answer) > 0:
                    continue
            if answer > 0 and answer.denominator <=30 and answer.numerator <= 30:
                questions.append(question)
                answers.append(answer)
                break
        paragraph = doc_out.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run(text=str(x+1) + "、")
        run.font.name = 'Cambria Math'
        run.font.size = Pt(14)
        latex_str = str(question[0]) + question[1] + str(question[2]) + " ="
        # 使用Word的OMath对象（需LaTeX转UnicodeMath）
        equation = f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">\
                    <m:oMath><m:r><m:t>{latex_str}</m:t></m:r></m:oMath></m:oMathPara>'
        run._element.append(parse_xml(equation))
    # 新建文件夹
    os.makedirs(f"./{folder}")
    # 生成二维码
    img = qrcode.make(', '.join(map(str, answers)))
    img.save(f"./{folder}/answer.png")
    doc_out.add_picture(f"./{folder}/answer.png", width=Cm(8))
    # 输出保存
    doc_out.save(f'./{folder}/questions.docx')
    print(questions)
    print(answers)
    answers_txt = open(f"./{folder}/answers.txt", "w")
    answers_txt.write(', '.join(map(str, answers)))
    answers_txt.close()

# 输入题型
while True:
    try:
        question_type = int(input("请选择题型：\n0、混合 1、加法 2、减法 3、乘法 4、除法（默认混合）：") or "0")
        if question_type in [0, 1, 2, 3, 4]:
            break
    except Exception as e:
        print("错误：" + str(e))
# 输入题量
while True:
    try:
        quantity = int(input("请输入题量（默认30）：") or "30")
        if quantity > 0:
            break
    except Exception as e:
        print("错误：" + str(e))
# 输入份数
while True:
    try:
        paper = int(input("请输入份数（默认1）：") or "1")
        if paper > 0:
            break
    except Exception as e:
        print("错误：" + str(e))
for p in range(paper):
    generate_paper(question_type, quantity, f"{question_types[question_type]} {datetime.now().strftime('%Y-%m-%d %H-%M-%S')} {p+1}")
