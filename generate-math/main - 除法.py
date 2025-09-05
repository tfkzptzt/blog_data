from docx import Document
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.shared import Pt
from fractions import Fraction
from random import randint, choice
import qrcode
# 创建变量
questions = []
answers = []
# 生成随机分数
def Random_Fraction():
    while True:
        f = Fraction(randint(1,10), randint(1,10))
        if f < 2 and f != 1:
            return f
# 生成随机问题
def Random_Question():
    f1 = Random_Fraction()
    # symbol=choice([" + ", " - ", " × ", " ÷ "])
    symbol = " ÷ "
    f2 = Random_Fraction()
    match symbol:
        case " + ":
            ans = f1 + f2
        case " - ":
            ans = f1 - f2
        case " × ":
            ans = f1 * f2
        case " ÷ ":
            ans = f1 / f2
    #ans = eval("f1" + symbol + "f2")
    return [[f1 , symbol , f2], ans]
# 输入题量
while True:
    try:
        quantity = int(input("请输入题量："))
        if quantity > 0:
            break
    except Exception as e:
        print("错误：" + str(e))
# 新建文档
doc_out = Document()
# 获取文档的第一个节（默认只有一个节）
section = doc_out.sections[0]
section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
# 设置页边距（单位为厘米）
section.top_margin = Cm(1.27) # 上边距
section.bottom_margin = Cm(1.27) # 下边距
section.left_margin = Cm(1.27) # 左边距
section.right_margin = Cm(1.27) # 右边距
paragraph = doc_out.add_paragraph()
paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = paragraph.add_run(text=str(quantity) + "道分数除法运算")
run.bold = True
run.font.name = '宋体'
run.font.size = Pt(24)
run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
paragraph = doc_out.add_paragraph()
paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = paragraph.add_run(text="班级：____________姓名：____________")
run.bold = True
run.font.name = '宋体'
run.font.size = Pt(12)
run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 开始生成
for x in range(quantity):
    while True:
        t = Random_Question()
        question = t[0]
        answer = t[1]
        if len(questions) > 0:
            if questions.count(answer) > 0:
                continue
        if answer > 0 and answer.denominator <=30 and answer.numerator <= 30:
            questions.append(question)
            answers.append(answer)
            break
    paragraph = doc_out.add_paragraph()
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(text=str(x+1) + "、")
    run.font.name = 'Cambria Math'
    run.font.size = Pt(14)
    latex_str = str(question[0]) + question[1] + str(question[2]) + " ="
    # 使用Word的OMath对象（需LaTeX转UnicodeMath）
    equation = f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">\
                <m:oMath><m:r><m:t>{latex_str}</m:t></m:r></m:oMath></m:oMathPara>'
    run._element.append(parse_xml(equation))
# 生成二维码
img = qrcode.make(', '.join(map(str, answers)))
img.save("./answer.png")
doc_out.add_picture("./answer.png", width=Cm(8)).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# 输出保存
doc_out.save('./sample.docx')
print(questions)
print(answers)
answers_txt = open("./answers.txt", "w")
answers_txt.write(', '.join(map(str, answers)))
answers_txt.close()
